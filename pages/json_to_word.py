import streamlit as st
from docx import Document
import json
from io import BytesIO

def convert_json_to_word(json_data):
    # Create a Word document
    doc = Document()
    doc.add_heading("Quiz Questions", level=1)

    for key, value in json_data.items():
        doc.add_heading(f"Question {value['no']}", level=2)
        doc.add_paragraph(value['mcq'])
        doc.add_paragraph("Options:")
        for opt_key, opt_value in value['options'].items():
            doc.add_paragraph(f"  {opt_key.upper()}: {opt_value}")
        doc.add_paragraph(f"Correct Answer: ____________")

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# def main():
#     st.title("JSON to Word Converter")
#     st.write("Upload a JSON file containing quiz data, or use a preloaded JSON file to create a Word document.")

#     # Directly initialize a JSON file for testing
#     uploaded_file = "quiz_response.json"  # Example file name

#     try:
#         # Open and read the JSON file
#         with open(uploaded_file, "r") as file:
#             json_data = json.load(file)

#         st.write("### JSON Preview")
#         st.json(json_data)

#         # Convert JSON to Word
#         word_file = convert_json_to_word(json_data)

#         # Create a download button
#         st.download_button(
#             label="Download Word Document",
#             data=word_file,
#             file_name="quiz_questions.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )

#     except FileNotFoundError:
#         st.error(f"File {uploaded_file} not found. Please upload a valid file or check the path.")
#     except json.JSONDecodeError:
#         st.error(f"File {uploaded_file} is not a valid JSON file. Please check the file content.")

# if __name__ == "__main__":
#     main()
