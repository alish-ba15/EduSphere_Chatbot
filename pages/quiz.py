import json
from dotenv import load_dotenv
# from langchain.llms import OpenAI
from langchain_google_genai import GoogleGenerativeAI
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.chains import SequentialChain
import streamlit as st
import traceback
import pandas as pd
from pages.pdf_design import load_quiz_data
# from langchain.callbacks import get_openai_callback
from pages.pdf_load import parse_file, get_table_data, RESPONSE_JSON
from docx import Document
import json
from io import BytesIO
import os

load_dotenv()

from menu import menu_with_redirect

# Redirect to app.py if not logged in, otherwise show the navigation menu
menu_with_redirect()




# This is an LLMChain to create 10-20 multiple choice questions from a given piece of text.
api_key = os.getenv("GOOGLE_API_KEY")
llm = GoogleGenerativeAI(model="gemini-pro")
template = """
Text: {text}
You are an expert MCQ maker. Given the above text, it is your job to\
create a quiz of {number} multiple choice questions for  students in {tone} tone.
Make sure that questions are not repeated and check all the questions to be conforming to the text as well.
Make sure to format your response like the RESPONSE_JSON below and use it as a guide.\
Ensure to make the {number} MCQs.
### RESPONSE_JSON
{response_json}
"""
quiz_generation_prompt = PromptTemplate(
    input_variables=["text", "number", "tone", "response_json"],
    template=template,
)
quiz_chain = LLMChain(
    llm=llm, prompt=quiz_generation_prompt, output_key="quiz", verbose=True
)

# This is an LLMChain to evaluate the multiple choice questions created by the above chain
# llm = OpenAI(model_name="gpt-3.5-turbo", temperature=0)
template = """You are an expert english grammarian and writer. Given a multiple choice quiz for  students.\
You need to evaluate complexity of the questions and give a complete analysis of the quiz if the students 
will be able to understand the questions and answer them. Only use at max 50 words for complexity analysis.
If quiz is not at par with the cognitive and analytical abilities of the students,\
update the quiz questions which need to be changed and change the tone such that it perfectly fits the students abilities. 
Quiz MCQs:
{quiz}
Critique from an expert english writer of the above quiz:"""

quiz_evaluation_prompt = PromptTemplate(
    input_variables=["quiz"], template=template
)
review_chain = LLMChain(
    llm=llm, prompt=quiz_evaluation_prompt, output_key="review", verbose=True
)

# This is the overall chain where we run these two chains in sequence.
generate_evaluate_chain = SequentialChain(
    chains=[quiz_chain, review_chain],
    input_variables=["text", "number","tone", "response_json"],
    # Here we return multiple variables
    output_variables=["quiz", "review"],
    verbose=True,
)

def convert_json_to_word(json_data):
    # Create a Word document
    doc = Document()
    doc.add_heading("Quiz Questions", level=1)

    for key, value in json_data.items():
        doc.add_heading(f"Question {value['no']}", level=2)
        doc.add_paragraph(value['mcq'])
        doc.add_paragraph("Options:")
        for opt_key, opt_value in value['options'].items():
            doc.add_paragraph(f"  {opt_key.upper()}: {opt_value}")
        doc.add_paragraph(f"Correct Answer: ____________")

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer




def main():
    st.title("Quiz Generator...!")

    # Create a form using st.form
    with st.form("user_inputs"):
        # File upload
        uploaded_file = st.file_uploader("Upload a PDF or text file")

        # Input fields
        mcq_count = st.number_input("No of MCQs", min_value=3, max_value=20)
        # grade = st.number_input("Insert Grade", min_value=1, max_value=10)
        tone = st.text_input("Insert Quiz tone", max_chars=100, placeholder="simple")

        button = st.form_submit_button("Create quiz")

    # Check if the button is clicked and all fields have inputs
    if button and uploaded_file is not None and mcq_count and tone:
        with st.spinner("Loading..."):
            try:
                text = parse_file(uploaded_file)

                # Count tokens and cost of API call
                response = generate_evaluate_chain(
                    {
                        "text": text,
                        "number": mcq_count,
                        "tone": tone,
                        "response_json": json.dumps(RESPONSE_JSON),
                    }
                )
            except Exception as e:
                traceback.print_exception(type(e), e, e.__traceback__)
                st.error("Error")
            else:
                print("Total Tokens")
                # print(f"Prompt Tokens: {cb.prompt_tokens}")
                # print(f"Completion Tokens: {cb.completion_tokens}")
                # print(f"Total Cost (USD): ${cb.total_cost}")

                if isinstance(response, dict):
                    # Extract quiz data from the response
                    quiz = response.get("quiz", None)
                    print(f"OHHH: {quiz}")
                    if quiz:
                        if "### RESPONSE_JSON" in quiz:
                            quiz = quiz.replace("### RESPONSE_JSON", "").strip()
                            try:
                                if not quiz.strip():
                                    raise ValueError("Quiz data is empty.")
                                quiz_data = json.loads(quiz)
                            except json.JSONDecodeError as e:
                                st.error(f"Failed to parse quiz JSON: {e}")
                            except ValueError as e:
                                st.error(f"Invalid quiz data: {e}")
                            else:
                                quiz_file = "quiz_response.json"
                                with open(quiz_file, "w") as file:
                                    json.dump(quiz_data, file, indent=4)
                                    st.success(f"Quiz saved to {os.path.abspath(quiz_file)}")

                        table_data = get_table_data(quiz)
                        if table_data is not None:
                            df = pd.DataFrame(table_data)
                            df.index = df.index + 1
                            st.table(df)
                            # Display the review in a text box
                            st.text_area(label="Review", value=response["review"])
                        else:
                            st.error("Error in table data")
                else:
                    st.write(response)

                try:
                    data = load_quiz_data("quiz_response.json")
                except Exception as e:
                    st.error(f"Could not load quiz data: {e}")
                    return

                try:
                    # Open and read the JSON file
                    with open("quiz_response.json", "r") as file:
                        json_data = json.load(file)

                
                    # Convert JSON to Word
                    word_file = convert_json_to_word(json_data)

                    # Create a download button
                    st.download_button(
                        label="Download Word Document",
                        data=word_file,
                        file_name="quiz_questions.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except FileNotFoundError:
                    st.error(f"File 'quiz_response.json' not found. Please upload a valid file or check the path.")
                except json.JSONDecodeError:
                    st.error("The file 'quiz_response.json' is not a valid JSON file. Please check the file content.")

                


if __name__ == "__main__":
    main()